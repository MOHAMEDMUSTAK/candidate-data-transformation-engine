from __future__ import annotations

"""
Eightfold AI - DOCX Resume Extractor

Extracts candidate information from DOCX resumes using python-docx.
Shares heuristic extraction logic with the PDF extractor.
"""

import logging
import io
from typing import Any

import docx

from backend.models.pipeline import IntermediateRecord
from backend.models.source import SourceType
from backend.extractors.pdf_extractor import (
    EMAIL_PATTERN,
    PHONE_PATTERN,
    LINKEDIN_PATTERN,
    GITHUB_PATTERN,
    URL_PATTERN,
    _extract_name,
    _extract_skills_from_text,
    _extract_section,
    _parse_experience_section,
    _parse_education_section,
)

logger = logging.getLogger(__name__)


def extract_docx(content: bytes, filename: str) -> list[IntermediateRecord]:
    """
    Extract candidate information from a DOCX resume.

    Args:
        content: Raw DOCX file bytes
        filename: Original filename

    Returns:
        List containing one IntermediateRecord
    """
    records: list[IntermediateRecord] = []

    try:
        doc = docx.Document(io.BytesIO(content))
    except Exception as e:
        logger.warning(f"DOCX '{filename}' could not be opened: {e}")
        return records

    # Extract all text from paragraphs
    text_parts: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text.strip())

    # Also try to extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text_parts.append(cell.text.strip())

    if not text_parts:
        logger.warning(f"DOCX '{filename}' contains no extractable text")
        return records

    full_text = "\n".join(text_parts)

    # Re-use the heuristic extraction logic from the PDF extractor
    name = _extract_name(full_text)
    emails = list(set(EMAIL_PATTERN.findall(full_text)))
    phones = list(set(PHONE_PATTERN.findall(full_text)))
    skills = _extract_skills_from_text(full_text)

    # Extract URLs
    linkedin = None
    github = None
    other_links: list[str] = []

    linkedin_matches = LINKEDIN_PATTERN.findall(full_text)
    if linkedin_matches:
        linkedin = linkedin_matches[0]

    github_matches = GITHUB_PATTERN.findall(full_text)
    if github_matches:
        github = github_matches[0]

    all_urls = URL_PATTERN.findall(full_text)
    for url in all_urls:
        if "linkedin" not in url.lower() and "github" not in url.lower():
            other_links.append(url)

    # Extract sections
    experience = _parse_experience_section(_extract_section(full_text, "experience"))
    education = _parse_education_section(_extract_section(full_text, "education"))

    # Clean phone numbers
    import re
    phones = [p.strip() for p in phones if 7 <= len(re.sub(r'\D', '', p)) <= 15]

    record = IntermediateRecord(
        source_type=SourceType.RESUME_DOCX,
        source_file=filename,
        full_name=name,
        emails=sorted(set(emails)),
        phones=sorted(set(phones)),
        skills=skills,
        experience=experience,
        education=education,
        linkedin=linkedin,
        github=github,
        other_links=other_links,
        raw_text=full_text[:5000],
    )
    records.append(record)

    logger.info(
        f"DOCX '{filename}': extracted name='{name}', "
        f"{len(emails)} emails, {len(phones)} phones, "
        f"{len(skills)} skills"
    )
    return records
